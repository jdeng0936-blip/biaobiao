"""
DOCX 导出服务 — 单元测试（pytest 格式）
验证 BidExportService.export_document 生成的字节流为合法 docx。
无磁盘写入，全部在内存中完成。
"""
import io
import unittest

try:
    from app.services.export_service import BidExportService
    HAS_EXPORT = True
except ImportError:
    HAS_EXPORT = False

SKIP_MSG = "export_service 或 python-docx 未安装"


@unittest.skipUnless(HAS_EXPORT, SKIP_MSG)
class TestBidExportService(unittest.TestCase):
    """BidExportService.export_document 测试"""

    @classmethod
    def setUpClass(cls):
        cls.service = BidExportService()
        cls.project_name = "南京奥体中心维修及附属设施改造工程"
        cls.sections = {
            "一、工程概况与整体部署": (
                "1.1 工程概况\n"
                "本工程位于南京市建邺区，涉及场馆防水维修。\n\n"
            ),
            "二、主要施工方案": (
                "2.1 防水工程施工方案\n"
                "采用 4mm 厚 SBS 改性沥青防水卷材。\n"
            ),
        }

    def test_returns_bytes(self):
        """导出结果为 bytes 类型"""
        result = self.service.export_document(self.project_name, self.sections)
        self.assertIsInstance(result, bytes)

    def test_non_empty_output(self):
        """导出结果非空"""
        result = self.service.export_document(self.project_name, self.sections)
        self.assertGreater(len(result), 100, "DOCX 字节流过小")

    def test_valid_docx_magic_bytes(self):
        """字节流以 PK (ZIP) 文件头开始（docx 本质是 zip）"""
        result = self.service.export_document(self.project_name, self.sections)
        self.assertTrue(result[:2] == b"PK", "不是合法的 ZIP/DOCX 文件头")

    def test_can_open_as_docx(self):
        """生成的字节流可被 python-docx 解析"""
        from docx import Document
        result = self.service.export_document(self.project_name, self.sections)
        doc = Document(io.BytesIO(result))
        # 至少包含段落
        self.assertGreater(len(doc.paragraphs), 0)

    def test_contains_project_name(self):
        """文档内容包含项目名称"""
        from docx import Document
        result = self.service.export_document(self.project_name, self.sections)
        doc = Document(io.BytesIO(result))
        full_text = "\n".join(p.text for p in doc.paragraphs)
        self.assertIn("南京奥体中心", full_text)

    def test_contains_section_titles(self):
        """文档内容包含章节标题"""
        from docx import Document
        result = self.service.export_document(self.project_name, self.sections)
        doc = Document(io.BytesIO(result))
        full_text = "\n".join(p.text for p in doc.paragraphs)
        self.assertIn("工程概况", full_text)
        self.assertIn("施工方案", full_text)

    def test_empty_sections(self):
        """空章节不崩溃"""
        result = self.service.export_document("空项目", {})
        self.assertIsInstance(result, bytes)
        self.assertGreater(len(result), 0)

    def test_with_company_name(self):
        """含公司名称参数不崩溃"""
        result = self.service.export_document(
            self.project_name, self.sections, company_name="江苏某某建筑集团"
        )
        self.assertIsInstance(result, bytes)


if __name__ == "__main__":
    unittest.main()
