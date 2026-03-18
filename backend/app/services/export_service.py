"""
标书 Word 导出引擎
生成符合 G 端政企标准的 .docx 文件

支持：
  - 标书封面
  - 目录页（自动编号）
  - 章节内容（中文编号层级）
  - 页眉页脚（项目名 + 页码）
  - 段落缩进（首行空两格）
"""
import io
import logging
import re
from typing import Optional

from docx import Document
from docx.shared import Pt, Cm, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn

logger = logging.getLogger("export_service")

# 中文编号层级样式映射
_HEADING_LEVELS = {
    1: {"prefix_pattern": r"^[一二三四五六七八九十]+、", "font_size": 16, "bold": True},
    2: {"prefix_pattern": r"^（[一二三四五六七八九十]+）", "font_size": 14, "bold": True},
    3: {"prefix_pattern": r"^\d+\.", "font_size": 13, "bold": True},
    4: {"prefix_pattern": r"^（\d+）", "font_size": 12, "bold": False},
    5: {"prefix_pattern": r"^[①②③④⑤⑥⑦⑧⑨⑩]", "font_size": 12, "bold": False},
}


class BidExportService:
    """标书 Word 导出服务"""

    def __init__(self):
        pass

    def _set_default_font(self, doc: Document, font_name: str = "仿宋_GB2312"):
        """设置文档默认字体"""
        style = doc.styles["Normal"]
        font = style.font
        font.name = font_name
        font.size = Pt(12)
        # 中文字体设置
        style.element.rPr.rFonts.set(qn("w:eastAsia"), font_name)

    def _add_header_footer(self, doc: Document, project_name: str):
        """添加页眉页脚"""
        for section in doc.sections:
            # 页眉：项目名
            header = section.header
            header_para = header.paragraphs[0]
            header_para.text = project_name
            header_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            header_run = header_para.runs[0] if header_para.runs else header_para.add_run()
            header_run.font.size = Pt(9)
            header_run.font.color.rgb = RGBColor(128, 128, 128)

            # 页脚：页码
            footer = section.footer
            footer_para = footer.paragraphs[0]
            footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            # 插入自动页码域
            run = footer_para.add_run()
            fldChar1 = run._element.makeelement(qn("w:fldChar"), {qn("w:fldCharType"): "begin"})
            run._element.append(fldChar1)
            run2 = footer_para.add_run()
            instrText = run2._element.makeelement(qn("w:instrText"), {})
            instrText.text = " PAGE "
            run2._element.append(instrText)
            run3 = footer_para.add_run()
            fldChar2 = run3._element.makeelement(qn("w:fldChar"), {qn("w:fldCharType"): "end"})
            run3._element.append(fldChar2)

    def _add_cover_page(self, doc: Document, project_name: str, company_name: str = ""):
        """添加封面"""
        # 空行占位
        for _ in range(6):
            doc.add_paragraph("")

        # 项目名称
        title = doc.add_paragraph()
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = title.add_run(project_name)
        run.font.size = Pt(26)
        run.font.bold = True
        run.font.name = "黑体"
        run.element.rPr.rFonts.set(qn("w:eastAsia"), "黑体")

        doc.add_paragraph("")

        # 副标题
        subtitle = doc.add_paragraph()
        subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
        sub_run = subtitle.add_run("投 标 文 件")
        sub_run.font.size = Pt(22)
        sub_run.font.bold = True

        doc.add_paragraph("")
        doc.add_paragraph("")

        # 公司名
        if company_name:
            company = doc.add_paragraph()
            company.alignment = WD_ALIGN_PARAGRAPH.CENTER
            c_run = company.add_run(f"投标人：{company_name}")
            c_run.font.size = Pt(16)

        # 分页
        doc.add_page_break()

    def _detect_heading_level(self, line: str) -> Optional[int]:
        """检测行是否为标题（基于中文编号模式）"""
        stripped = line.strip()
        for level, config in _HEADING_LEVELS.items():
            if re.match(config["prefix_pattern"], stripped):
                return level
        return None

    def _add_content_paragraph(self, doc: Document, text: str, is_heading: bool = False, level: int = 0):
        """添加内容段落"""
        para = doc.add_paragraph()

        if is_heading and level > 0:
            config = _HEADING_LEVELS.get(level, {})
            run = para.add_run(text)
            run.font.size = Pt(config.get("font_size", 12))
            run.font.bold = config.get("bold", False)
            run.font.name = "黑体"
            run.element.rPr.rFonts.set(qn("w:eastAsia"), "黑体")

            # 标题前后间距
            para.paragraph_format.space_before = Pt(12)
            para.paragraph_format.space_after = Pt(6)
        else:
            # 正文段落：首行缩进两个汉字
            run = para.add_run(text)
            run.font.size = Pt(12)
            run.font.name = "仿宋_GB2312"
            run.element.rPr.rFonts.set(qn("w:eastAsia"), "仿宋_GB2312")
            para.paragraph_format.first_line_indent = Cm(0.74)  # 约两个汉字宽
            para.paragraph_format.line_spacing = Pt(28)  # 28磅行距（标书常用）

    def export_document(
        self,
        project_name: str,
        sections: dict[str, str],
        company_name: str = "",
    ) -> bytes:
        """
        导出标书为 Word 文档

        参数:
            project_name: 项目名称
            sections: 章节内容 {"章节标题": "章节正文"}
            company_name: 投标公司名

        返回:
            .docx 文件的字节内容
        """
        doc = Document()

        # 页面设置 (A4)
        section = doc.sections[0]
        section.page_width = Cm(21.0)
        section.page_height = Cm(29.7)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.5)
        section.top_margin = Cm(2.54)
        section.bottom_margin = Cm(2.54)

        # 默认字体
        self._set_default_font(doc)

        # 页眉页脚
        self._add_header_footer(doc, project_name)

        # 封面
        self._add_cover_page(doc, project_name, company_name)

        # 目录页占位（Word 目录需手动更新域）
        toc_title = doc.add_paragraph()
        toc_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        toc_run = toc_title.add_run("目  录")
        toc_run.font.size = Pt(18)
        toc_run.font.bold = True
        doc.add_paragraph("(生成后请在 Word 中右键目录 → 更新域)")
        doc.add_page_break()

        # 逐章节写入
        for title, content in sections.items():
            # 章节标题
            chapter_para = doc.add_paragraph()
            chapter_run = chapter_para.add_run(title)
            chapter_run.font.size = Pt(16)
            chapter_run.font.bold = True
            chapter_run.font.name = "黑体"
            chapter_run.element.rPr.rFonts.set(qn("w:eastAsia"), "黑体")
            chapter_para.paragraph_format.space_before = Pt(24)
            chapter_para.paragraph_format.space_after = Pt(12)

            # 章节内容按行解析
            lines = content.split("\n")
            for line in lines:
                stripped = line.strip()
                if not stripped:
                    continue

                heading_level = self._detect_heading_level(stripped)
                if heading_level:
                    self._add_content_paragraph(doc, stripped, is_heading=True, level=heading_level)
                else:
                    self._add_content_paragraph(doc, stripped)

            # 章节间分页
            doc.add_page_break()

        # 导出为字节
        buffer = io.BytesIO()
        doc.save(buffer)
        buffer.seek(0)
        return buffer.getvalue()
